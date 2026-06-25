# -*- coding: utf-8 -*-
"""
《算法过程可视化系统》需求分析文档（规范版）

格式严格遵循《软件工程专业毕业设计（论文）写作规范与要求(2026)》：
- 页边距：上 3.3 cm，下 2.3 cm，左 2.8 cm，右 2.3 cm
- 正文：宋体小四号(12pt)，行距 1.5 倍，首行缩进 2 字符（西文 Times New Roman）
- 论文题目：二号黑体(22pt)
- 一级标题（章）：小三号黑体(15pt)
- 二级标题（节）：四号黑体(14pt)
- 三级标题（条）：小四号黑体(12pt)
- 四级标题（款）：小四号黑体(12pt)
- 页码：小五号(9pt) 底端居中

内容根据实际产品截图（localhost:5173/graph/dijkstra，Vue.js + Vite，深色主题）
和成员分工表（Vue.js 3 + Vite + Canvas，4 人 2 周开发周期）调整。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== 字号常量（规范定义） =====
SZ_ER_HAO       = 22   # 二号
SZ_XIAO_ER_HAO  = 18   # 小二号
SZ_SAN_HAO      = 16   # 三号
SZ_XIAO_SAN_HAO = 15   # 小三号
SZ_SI_HAO       = 14   # 四号
SZ_XIAO_SI_HAO  = 12   # 小四号
SZ_WU_HAO       = 10.5 # 五号
SZ_XIAO_WU_HAO  = 9    # 小五号


def set_run_font(run, cn_name="宋体", en_name="Times New Roman",
                 size=SZ_XIAO_SI_HAO, bold=False, color=None):
    """按规范设置中文/西文字体。中文由 eastAsia 控制，西文由 ascii/hAnsi 控制。"""
    run.font.name = en_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'), en_name)
    rFonts.set(qn('w:hAnsi'), en_name)


def add_body(doc, text, size=SZ_XIAO_SI_HAO, bold=False,
             align=None, first_line_indent=True, line_spacing=1.5):
    """正文段：宋体小四号，1.5 倍行距，首行缩进 2 字符。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", size=size, bold=bold)
    return p


def add_h1(doc, text):
    """一级标题（章）：小三号黑体，居中，段前 0.5 行段后 0.5 行。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(8)
    pf.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_XIAO_SAN_HAO, bold=True)
    return p


def add_h2(doc, text):
    """二级标题（节）：四号黑体，左顶格。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_SI_HAO, bold=True)
    return p


def add_h3(doc, text):
    """三级标题（条）：小四号黑体，左顶格。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_XIAO_SI_HAO, bold=True)
    return p


def add_h4(doc, text):
    """四级标题（款）：小四号黑体，左顶格 ，与正文字体同号但加粗。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_XIAO_SI_HAO, bold=True)
    return p


def add_bullet(doc, text, size=SZ_XIAO_SI_HAO):
    """列表项：宋体小四号，左缩进 2 字符。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Pt(size * 2)
    pf.first_line_indent = Pt(0)
    run = p.add_run("（" + str(getattr(add_bullet, "_n", 1)) + "）" + text) \
        if False else p.add_run("● " + text)
    set_run_font(run, "宋体", "Times New Roman", size=size)
    return p


def add_caption(doc, text):
    """图/表标题：五号黑体居中。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "黑体", "Times New Roman", size=SZ_WU_HAO, bold=True)
    return p


def add_code_block(doc, text):
    """等宽字体代码块：Consolas 五号。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(24)
    pf.line_spacing = 1.2
    pf.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", "Consolas", size=SZ_WU_HAO)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """表格：表头黑体小四号居中；正文宋体五号。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", "Times New Roman",
                     size=SZ_WU_HAO, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.first_line_indent = Pt(0)
            run = p.add_run(str(val))
            set_run_font(run, "宋体", "Times New Roman",
                         size=SZ_WU_HAO)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_page_number(section):
    """规范要求：页码小五号底端居中。"""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run = p.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, "宋体", "Times New Roman",
                 size=SZ_XIAO_WU_HAO)


def main():
    doc = Document()

    # ============ 页面设置（规范） ============
    section = doc.sections[0]
    section.top_margin = Cm(3.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.3)
    add_page_number(section)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(SZ_XIAO_SI_HAO)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # ============ 封面 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("算法过程可视化系统")
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_ER_HAO, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("需求分析文档")
    set_run_font(run, "黑体", "Times New Roman",
                 size=SZ_ER_HAO, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    # 文档属性表
    headers = ["文档属性", "内容"]
    rows = [
        ["项目名称", "Algorithm Process Visualization System（算法过程可视化系统）"],
        ["应用名称", "算法过程可视化系统"],
        ["App ID", "cn.algoviz.web"],
        ["编制日期", "2026-06-19"],
        ["编制小组", "孟欣（组长）、李佳乐、宋启悦、冀文卓"],
        ["项目周期", "2026-06-19 ~ 2026-07-02（共 2 周）"],
        ["引用规范", "OMG UML 2.5.1、DBML 1.0、IEEE 830-1998、GB/T 8567-2006、"
                    "《软件工程专业毕业设计（论文）写作规范与要求(2026)》"],
    ]
    add_table(doc, headers, rows, col_widths=[3.5, 12.0])

    doc.add_page_break()

    # ============ 目录（手写） ============
    add_h1(doc, "目  录")
    toc_items = [
        "1  问题定义",
        "    1.1  业务/教学背景",
        "    1.2  核心问题",
        "    1.3  项目定位与边界",
        "2  可行性研究",
        "    2.1  技术可行性",
        "    2.2  经济可行性",
        "    2.3  操作可行性",
        "    2.4  法律/合规可行性",
        "    2.5  进度可行性",
        "3  需求分析",
        "    3.1  系统目标",
        "    3.2  用户群体",
        "    3.3  功能需求",
        "    3.4  非功能需求",
        "4  UML 与数据建模",
        "5  AI 辅助需求分析的应用",
        "6  附录",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(item)
        set_run_font(run, "宋体", "Times New Roman",
                     size=SZ_XIAO_SI_HAO)

    doc.add_page_break()

    # ============ 1. 问题定义 ============
    add_h1(doc, "1  问题定义")

    add_h2(doc, "1.1  业务/教学背景")
    add_body(doc,
        "数据结构与算法是计算机相关专业的核心课程，其学习成效直接决定后续操作系统、数据库、"
        "编译原理等课程的学习深度。然而在《数据结构》课程的实际教学中，长期存在以下痛点："
    )
    pain_points = [
        "抽象概念难以建立心智模型：链表的指针指向、栈/队列的入出顺序、二叉树的递归遍历、"
        "Dijkstra 的松弛过程等动态行为，仅凭伪代码与静态插图难以在脑海中复现。",
        "执行过程不可回溯：传统课堂演示一过即逝，学生在课后复习时无法逐步暂停、回放某一关键步骤；"
        "终端控制台的 printf 调试只能输出最终结果或离散日志，无法以图形方式呈现状态演变。",
        "代码与图形脱节：现有 PPT 动画往往与真实代码无关，学生看到的“动画”与教材中的代码片段"
        "无法对应，造成“看动画似乎懂了，写代码又不会”的割裂感。",
        "测试用例难以复现：手工构造图、矩阵、稀疏矩阵等结构耗时冗长，学生在做对比实验、调参实验时"
        "缺乏可保存、可复用的测试用例库。",
        "现有工具存在局限：在线可视化网站（如 visualgo、cs-visualizer）依赖网络且算法覆盖有限；"
        "桌面 IDE 调试器虽然精准但缺乏面向数据结构的可视化语义；二者皆不易二次开发与定制。",
    ]
    for pp in pain_points:
        add_bullet(doc, pp)

    add_h2(doc, "1.2  核心问题")
    add_body(doc,
        "如何为《数据结构》课程构建一个基于浏览器的、覆盖典型数据结构与算法的交互式实验平台，"
        "使每一行真实代码都能与可视化画布上的状态变化精确同步，并支持自定义测试数据与多档播放速度，"
        "从而真正打通“代码 — 状态 — 图形”的认知闭环？"
    )

    add_h2(doc, "1.3  项目定位与边界")
    add_body(doc,
        "本项目定位为《数据结构》课程设计的 Web 端算法可视化实验平台，"
        "以“专用画布 + 源码级同步 + 测试数据闭环”为三大核心能力，"
        "目标用户群体为高校教师与本科生；产品对外展示名称即为“算法过程可视化系统”。"
    )

    add_h3(doc, "1.3.1  项目范围（系统将提供）")
    scope = [
        "覆盖《数据结构》课程主流算法的可视化：包括线性表（顺序表/单链表/双链表/循环链表）、"
        "栈与队列（顺序表/链表/循环队列、括号匹配、表达式计算）、树结构（二叉树/哈夫曼树）、"
        "图结构（邻接矩阵/邻接表/Dijkstra 最短路径）以及典型递归与排序算法等。",
        "采用 Vue.js 3 + Vite + Canvas 技术栈构建单页 Web 应用，"
        "本地开发服务器监听 5173 端口，部署形态为静态资源，可托管于校内 Web 服务器或本地直接打开。",
        "发布形式包括 Vite 构建出的静态发行包（HTML + JS + CSS）以及开发期热重载环境，"
        "适合教师课堂演示、学生课后自学与课程作业实验。",
        "核心能力包括：算法分类目录、专用算法画布、源码级同步、测试用例选择、"
        "随机数据生成、播放/暂停/单步/重置控制、播放速度调节、当前步骤说明、复杂度展示。",
    ]
    for i, s in enumerate(scope, start=1):
        add_bullet(doc, s)

    add_h3(doc, "1.3.2  边界划定原则")
    bounds = [
        "面向单机使用，不引入登录、云同步、社交分享等需要联网的能力。",
        "不包含算法源码的在线编译/执行能力，平台内置算法实现由开发团队预置，"
        "用户不可任意上传执行任意代码。",
        "不包含课程管理（作业发布、批改、评分）功能，平台聚焦“可视化实验”而非“教学管理”。",
        "不引入用户体系，不收集任何个人数据；用户输入的测试数据仅在浏览器会话中保留。",
        "首版不支持移动端深度适配，动画呈现质量与交互体验以桌面分辨率（≥ 1366×768）为基准。",
    ]
    for b in bounds:
        add_bullet(doc, b)

    # ============ 2. 可行性研究 ============
    add_h1(doc, "2  可行性研究")

    add_h2(doc, "2.1  技术可行性")
    rows = [
        ["前端框架", "Vue.js 3.4 + Composition API",
         "成熟稳定，国内高校与企业广泛采用；响应式系统天然契合“数据 → 视图”同步场景；"
         "学习曲线平缓，团队成员熟悉，风险低。"],
        ["构建工具", "Vite 5",
         "基于原生 ES Modules，开发期热重载快，生产期产物体积小；"
         "默认监听 5173 端口；与 Vue 生态无缝整合。"],
        ["UI 与样式", "原生 CSS + 定制深色主题",
         "采用统一深色主题（深蓝 #0F172A 系列），减少第三方依赖；"
         "通过 CSS 变量管理主题色，便于后续扩展浅色主题。"],
        ["可视化渲染", "HTML5 Canvas 2D / SVG",
         "Canvas 适合密集帧动画（如排序柱状图、节点遍历）；"
         "SVG 适合可点击的图节点；二者按场景互补，浏览器原生支持，零依赖。"],
        ["路由", "Vue Router 4",
         "为不同算法分配独立路由（如 /graph/dijkstra、/sort/quick），"
         "便于地址直达与状态隔离。"],
        ["状态管理", "Pinia（或组合式 store）",
         "管理“当前算法 / 当前步骤 / 播放状态 / 测试用例 ID”等全局状态；"
         "API 简洁，调试方便。"],
        ["数据结构", "纯 JavaScript 对象",
         "图、链表、树等结构均以 ES 对象表示，便于序列化保存与复现；"
         "无需引入额外库。"],
        ["开发工具链", "Node.js 18 + npm + ESLint + Prettier",
         "团队采用统一格式化规则与 Lint，保证代码风格一致；"
         "成员均具备 Vue.js 项目开发经验。"],
        ["测试", "Vitest + 手工 E2E",
         "Vitest 与 Vite 同源，零配置；E2E 阶段以课堂演示用例进行手工回归。"],
    ]
    add_table(doc, ["关键技术", "拟选型", "论证说明"],
              rows, col_widths=[2.6, 4.0, 8.9])
    add_body(doc,
        "结论：所选技术栈整体处于成熟稳定阶段，社区文档完整，团队成员均有相关经验，"
        "技术风险等级评定为 低；在 2 周开发周期内具备完整的技术落地条件。"
    )

    add_h2(doc, "2.2  经济可行性")
    econ = [
        "授权成本：所选 Vue.js、Vite、Vue Router、Pinia、ESLint、Prettier 等均采用 MIT/BSD/Apache-2.0 "
        "等宽松开源协议，零授权成本。",
        "运维成本：Web 端应用，构建产物为纯静态文件，可托管于校内已有 Web 服务器或本地文件协议直接打开，"
        "几乎零运维成本。",
        "硬件成本：开发机为团队成员自有设备，目标运行环境对硬件无苛刻要求（4 GB 内存即可流畅运行）。",
        "人力成本：4 名学生开发者在课程周期内并行投入，无额外人力支出，符合课程设计教学经费预算。",
        "潜在收益：① 完成《数据结构》课程设计任务并获得课程学分；② 形成可对外开源的教学辅助工具，"
        "提升团队工程化能力与作品集；③ 平台可二次扩展为更多算法的实验环境，长期复用价值高。",
    ]
    for e in econ:
        add_bullet(doc, e)
    add_body(doc, "综上所述，项目投入产出比合理，经济可行性 充分。")

    add_h2(doc, "2.3  操作可行性")
    ops = [
        "上手难度低：用户从顶部标签或左侧目录选择算法 → 在数据区填写或随机生成数据 → 点击“使用输入” → "
        "点击“播放”即可观看动画，整个核心闭环不超过 4 步操作。",
        "功能入口清晰：顶部为算法切换标签栏，左侧为输入数据 / 测试用例 / 播放控制 / 当前步骤 / 复杂度 "
        "五个面板，中央为可视化画布，按“信息密度”分区组织，避免单页过载。",
        "输入容错完善：输入框对非法字符、长度越界、邻接矩阵非方阵、起点/终点不存在等情况进行校验，"
        "并以红色边框 + 文案提示，立即反馈，不会因为脏数据导致动画崩溃。",
        "辅助功能丰富：随机图按钮、3 条预置测试用例（5/6/4 节点图）、播放速度滑块（约 700 ms 默认）、"
        "上一步/下一步、重置、当前步骤说明等共同降低使用门槛。",
        "学习曲线平滑：所有算法共享同一交互范式（选择 → 输入 → 播放 → 观察），"
        "用户掌握第一个算法后，其余算法的操作几乎零迁移成本。",
    ]
    for o in ops:
        add_bullet(doc, o)

    add_h2(doc, "2.4  法律/合规可行性")
    laws = [
        "开源协议兼容性：所选第三方依赖均采用 MIT、BSD、Apache-2.0 等宽松协议，彼此兼容；"
        "项目计划以 MIT License 发布。",
        "数据隐私合规：本项目不采集、不上传、不传输任何用户个人信息；用户输入的测试数据仅在浏览器内存中存在，"
        "刷新页面即清除，符合 GDPR、《个人信息保护法》（PIPL）的数据最小化原则。",
        "知识产权：所有原创算法实现、画布渲染逻辑、UI 设计均由项目组独立完成；"
        "引用第三方资源时在“关于”面板中明确署名；不抓取、不二次分发任何受保护内容。",
        "品牌商标：产品名“算法过程可视化系统”为通用描述性名称，不与任何已注册商标冲突；"
        "Logo 与字体均来自团队自绘或采用允许商用的开源字库（如思源黑体）。",
    ]
    for l in laws:
        add_bullet(doc, l)
    add_body(doc, "综上，项目在法律与合规层面无显著风险，可放心实施。")

    add_h2(doc, "2.5  进度可行性")
    sched_rows = [
        ["阶段 1 — 前期准备", "2026-06-19", "需求分析文档、可行性分析报告、技术选型报告、用例分析文档、成员分工表"],
        ["阶段 2 — 系统设计", "2026-06-21", "系统架构设计、UI 原型与组件结构、模块接口规范、数据流与状态机设计"],
        ["阶段 3 — 模块开发", "2026-06-26", "线性表 / 栈队列 / 树结构 / 图结构 / 排序 / 递归回溯 等可视化模块；交互组件与代码/步骤面板"],
        ["阶段 4 — 系统集成", "2026-06-28", "模块联调、构建打包、性能与浏览器兼容性测试"],
        ["阶段 5 — 文档整理", "2026-07-01", "用户手册、测试报告、系统实现文档、README、源码注释整理"],
        ["阶段 6 — 答辩准备", "2026-07-02", "答辩 PPT、演示视频、源码材料包、演示脚本与问答准备"],
    ]
    add_table(doc, ["阶段节点", "截止时间", "主要交付工作"],
              sched_rows, col_widths=[3.6, 2.6, 9.3])
    add_body(doc,
        "人力分工模式：组长孟欣负责系统架构、框架搭建与线性表模块；李佳乐负责栈队列模块与界面设计；"
        "宋启悦负责树结构与递归算法模块及测试体系；冀文卓负责交互组件、回溯算法模块与系统打包。"
        "并行推进策略：第 1 周第 1 天确定模块接口，使所有成员可并行开发；每日 20:00 例会同步进度并解决阻塞；"
        "组长在第 2 周第 1 天前完成主框架，使其余模块可即时集成。基于上述分工与并行策略，"
        "在 2 周周期内可按节点完成全部交付。"
    )

    # ============ 3. 需求分析 ============
    add_h1(doc, "3  需求分析")

    add_h2(doc, "3.1  系统目标")
    add_body(doc,
        "以下目标编号以 G 开头，每条目标配套可量化的验收标准。"
    )
    goal_rows = [
        ["G1", "覆盖《数据结构》课程主流算法",
         "至少实现线性表、栈队列、树结构、图结构、排序、递归回溯六大类共 20 种以上算法的可视化，"
         "每个算法均提供专用画布与代码同步。"],
        ["G2", "保证“代码 - 状态 - 画布”三方同步",
         "100% 算法支持代码行级高亮，每帧均能映射至唯一一行真实源码与一条自然语言步骤说明，"
         "且步骤序号严格单调（如 1/30、2/30 …）。"],
        ["G3", "提供完整的播放控制",
         "支持播放/暂停、单步前进、单步后退、重置；速度滑块（默认 700 ms，可在 100–2000 ms 之间调节）；"
         "在 i5-8500U 同级硬件上单步切换响应 ≤ 200 ms。"],
        ["G4", "测试数据闭环",
         "支持随机生成、手动输入、预置测试用例一键载入；每个算法不少于 3 条预置测试用例。"],
        ["G5", "复杂度与步骤说明",
         "每个算法在执行过程中实时显示时间复杂度（如 O((V+E) log V)）、空间复杂度（如 O(V)）"
         "以及当前步骤的自然语言说明。"],
        ["G6", "本地离线运行",
         "构建产物为纯静态文件，安装包 ≤ 50 MB；启动时间 ≤ 2 s（首次加载）；断网情况下所有功能正常工作。"],
        ["G7", "界面统一与无障碍",
         "全平台采用统一的深色主题；通过 WCAG 2.1 AA 级色彩对比；动画可随播放进度暂停以适配晕动症用户。"],
        ["G8", "可扩展架构",
         "新增一个算法仅需实现统一接口（init/step/reset）并注册到目录配置文件，框架代码零改动。"],
    ]
    add_table(doc, ["编号", "目标描述", "度量/验收方式"],
              goal_rows, col_widths=[1.6, 4.4, 9.5])

    add_h2(doc, "3.2  用户群体")
    add_body(doc,
        "本系统为 单机单用户 应用，不引入账号体系；所有偏好与数据均按浏览器会话进行隔离。"
        "下列角色用于功能场景刻画与用例建模。"
    )
    user_rows = [
        ["STU", "数据结构课程在读本科生（计算机/软件工程相关专业），编程基础初级到中级",
         "课前预习链表/树/图等抽象结构；课后复习老师讲解过的算法；做课程实验时验证自己手写代码逻辑。",
         "希望快速、直观地看到算法每一步的状态变化；能对自己的测试用例反复观察；学习曲线低。"],
        ["TCH", "高校《数据结构》课程任课教师或助教",
         "课堂演示、布置实验题；为学生准备多组对比示例；考前点对点辅导。",
         "演示画面美观、流畅；能保存预置课件用例；可投影到大屏不失真。"],
        ["DEV", "希望自学算法的在职开发者或考研学生",
         "工作之余复习、考研复试前突击；通过实验加深印象。",
         "完整覆盖经典算法；不依赖网络。"],
    ]
    add_table(doc, ["角色标识", "角色描述", "主要使用场景", "核心诉求"],
              user_rows, col_widths=[1.5, 4.0, 5.0, 5.0])

    add_h2(doc, "3.3  功能需求")
    add_body(doc, "功能需求编号遵循 IEEE 830 风格（FR-XXX），按功能模块成组组织。")

    add_h3(doc, "3.3.1  模块 1：导航与布局（NAV）")
    fr_nav = [
        ("FR-001", "顶部导航栏自左到右依次显示：应用 Logo + 名称（“算法过程可视化系统”）、"
                  "当前算法标签（如“冒泡排序 简单”、“快速排序 中等”、“Dijkstra 中等+”），"
                  "用户点击标签可快速切换不同算法。"),
        ("FR-002", "切换算法时主区域应即时切换至对应算法画布，并保留上一算法的输入历史；"
                  "切换过程必须重置当前播放状态，避免上一算法残留动画。"),
        ("FR-003", "系统采用统一深色主题（背景 #0F172A，主色 #3B82F6，强调色 #F59E0B）；"
                  "节点高亮、当前操作元素采用对比色提示（如黄色高亮当前源点）。"),
        ("FR-004", "页面布局采用左右两栏：左侧固定宽度的控制面板，右侧弹性宽度的可视化画布；"
                  "在窗口缩放时画布自适应。"),
    ]
    for fid, desc in fr_nav:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.2  模块 2：输入数据（DATA）")
    fr_data = [
        ("FR-010", "“输入数据”面板根据当前算法类型显示不同输入控件：图算法显示节点数、源点选择、"
                  "邻接边列表（每行格式“起点,终点,权重”）；排序算法显示数组输入框。"),
        ("FR-011", "用户可点击“随机图/随机数据”按钮，根据当前算法生成合规数据：图算法生成连通图，"
                  "数组生成指定长度的随机数组；点击“使用输入”按钮将当前输入应用到画布。"),
        ("FR-012", "面板提供至少 3 条预置“测试用例”（如 5 节点图 / 6 节点图 / 4 节点图），"
                  "用户点击即可一键载入对应数据。"),
        ("FR-013", "系统应对非法输入（非数字、起点/终点超出节点数、空输入等）进行校验，"
                  "以红色边框与文案提示用户。"),
    ]
    for fid, desc in fr_data:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.3  模块 3：算法可视化画布（CAN）")
    fr_can = [
        ("FR-020", "系统应为每一类数据结构提供专用画布渲染器（不套用统一模板）：图算法以节点+加权边、"
                  "线性表以节点+指针、栈/队列以纵/横向容器、树以分层节点、排序以柱状条带。"),
        ("FR-021", "画布上节点应显示节点标签（如 A、B、C）；边应显示权重；"
                  "图算法的每个节点下方显示当前最短距离（如 0、∞）。"),
        ("FR-022", "在算法每一步执行时，应高亮当前操作元素：当前源点、被访问节点、被松弛的边、"
                  "被比较元素等，至少使用三种以上颜色语义（默认 / 当前 / 已访问）。"),
        ("FR-023", "画布应在窗口尺寸变化时自适应缩放；元素过多时自动启用滚动条而非压缩元素。"),
    ]
    for fid, desc in fr_can:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.4  模块 4：播放控制（PLAY）")
    fr_play = [
        ("FR-030", "“播放控制”面板提供 4 个按钮：重置、上一步、播放、下一步。"
                  "下方提供速度滑块（默认 700 ms，可在 100–2000 ms 间调节）。"),
        ("FR-031", "面板下方实时显示当前进度（如“步骤 1 / 30”），并随播放推进自动更新。"),
        ("FR-032", "系统应实现状态机驱动：idle / running / paused / completed 四态可控转换；"
                  "已完成状态下禁用播放与单步按钮。"),
        ("FR-033", "播放控制操作应支持键盘快捷键：Space 播放/暂停、← / → 单步、R 重置。"),
    ]
    for fid, desc in fr_play:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.5  模块 5：当前步骤说明（STEP）")
    fr_step = [
        ("FR-040", "“当前步骤”面板用一段简短自然语言描述当前帧的核心操作，"
                  "例如“初始化：源点 A(0)，dist[0]=0，其余为 ∞”。"),
        ("FR-041", "步骤说明随播放进度实时刷新；用户可在“暂停”状态下停留观察。"),
        ("FR-042", "步骤说明应紧扣源码逻辑（如对应代码行的语义），便于学习者建立“代码-状态”映射。"),
    ]
    for fid, desc in fr_step:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.6  模块 6：复杂度展示（COMP）")
    fr_comp = [
        ("FR-050", "“复杂度”面板始终显示当前算法的时间复杂度（如 O((V+E) log V)）"
                  "与空间复杂度（如 O(V)）。"),
        ("FR-051", "复杂度信息以醒目颜色显示，便于学习者一眼对比不同算法的性能差异。"),
        ("FR-052", "面板支持鼠标悬浮显示复杂度推导简要说明（V 为节点数，E 为边数等）。"),
    ]
    for fid, desc in fr_comp:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h3(doc, "3.3.7  模块 7：算法目录与详情（NAV2）")
    fr_dir = [
        ("FR-060", "应用支持按分类（线性表 / 栈队列 / 树 / 图 / 排序 / 递归）展开算法目录，"
                  "用户可点击进入对应算法画布。"),
        ("FR-061", "每个算法应在画布顶部显示标题（如“Dijkstra 最短路径（中等+）”）和简要说明，"
                  "如“贪心算法，每次从未访问节点中选择距离源点最近的节点，然后松弛其所有邻接边”。"),
        ("FR-062", "算法名称旁的难度徽章采用文字（简单 / 中等 / 中等+ / 高），分别对应不同的颜色背景。"),
    ]
    for fid, desc in fr_dir:
        add_bullet(doc, "%s：%s" % (fid, desc))

    add_h2(doc, "3.4  非功能需求")

    add_h3(doc, "3.4.1  性能需求")
    perf = [
        "动画帧率：可视化画布在常规规模（数组长度 ≤ 100、图节点 ≤ 30、树高度 ≤ 8）下保持 ≥ 50 FPS。",
        "响应时间：单步切换、用例载入、面板切换响应时间 ≤ 200 ms（基于 i5-8500U / 8 GB 内存）。",
        "启动时间：首次加载 ≤ 2 s，热加载（已缓存）≤ 1 s。",
        "资源占用：浏览器进程内存 ≤ 200 MB，CPU 空闲态 ≤ 3%，播放态 ≤ 25%。",
        "构建产物：生产环境 JS 总大小 ≤ 1 MB（gzip 后），首屏请求数 ≤ 10。",
    ]
    for x in perf:
        add_bullet(doc, x)

    add_h3(doc, "3.4.2  兼容性需求")
    compat = [
        "操作系统：Windows 10/11、macOS 11+、Ubuntu 20.04+ 主流发行版。",
        "浏览器：Chrome 100+、Firefox 100+、Edge 100+；Safari 15+；建议使用 Chromium 内核浏览器。",
        "屏幕分辨率：最低支持 1366×768，推荐 1920×1080；支持高分辨率（≥ 200% 缩放）下文字与图形不失真。",
        "硬件要求：CPU x64 双核 1.6 GHz 以上、内存 4 GB 以上、磁盘剩余 ≥ 200 MB。",
        "GPU：要求支持 Canvas 2D 硬件加速，否则自动降级为低帧率渲染并提示用户。",
    ]
    for x in compat:
        add_bullet(doc, x)

    add_h3(doc, "3.4.3  易用性需求")
    usab = [
        "交互规范：遵循平台原生快捷键习惯（Ctrl/⌘ 通用），按钮命中区域 ≥ 32×32 px；"
        "表单组件遵循统一深色主题。",
        "学习成本：新用户从启动到完成第一次算法播放不超过 1 分钟；首屏内有完整的可视化提示。",
        "操作容错：所有破坏性操作（清空输入、重置）需弹窗二次确认；导入非法输入给出明确错误信息。",
        "国际化预留：UI 文案集中于资源字典，便于后续接入英文等其他语言。",
    ]
    for x in usab:
        add_bullet(doc, x)

    add_h3(doc, "3.4.4  安全与合规需求")
    sec = [
        "数据本地性：用户输入的测试数据仅在当前浏览器会话中存在，不通过网络外发。",
        "隐私采集规则：不采集用户个人信息、不上报埋点。",
        "依赖供应链：定期使用 npm audit 扫描第三方依赖漏洞，CVSS ≥ 7 的高危漏洞需在 1 周内升级。",
        "XSS 防护：用户输入的字符串在显示前一律转义，不直接通过 v-html 渲染。",
    ]
    for x in sec:
        add_bullet(doc, x)

    add_h3(doc, "3.4.5  无障碍需求")
    a11y = [
        "颜色对比：在深色主题下文字与背景对比度 ≥ 4.5:1，达到 WCAG 2.1 AA 级。",
        "色弱适配：高亮色不仅依赖颜色，亦伴随形状/边框变化（如选中节点同时加粗描边）。",
        "动效适配：用户可通过暂停按钮随时停止动画，便于晕动症用户。",
        "键盘可达：核心操作（播放/暂停、单步、切换算法）均可通过键盘完成，无需鼠标。",
    ]
    for x in a11y:
        add_bullet(doc, x)

    add_h3(doc, "3.4.6  可维护性需求")
    maint = [
        "组件复用规则：通用控件（按钮、滑块、面板）统一封装于 components/ui/ 目录；"
        "算法组件遵循统一接口契约 { name, difficulty, complexity, render(), runStep() }。",
        "代码规范：ESLint + Prettier 强制约束；每次提交前自动格式化。",
        "命名规范：组件 PascalCase（如 LinearList.vue）；函数 camelCase；事件 kebab-case。",
        "测试覆盖：核心算法逻辑单测覆盖率 ≥ 70%；关键交互手工 E2E 用例 ≥ 10 条。",
        "可扩展性：新增一个算法仅需实现统一接口并注册到目录配置文件，框架代码零改动。",
    ]
    for x in maint:
        add_bullet(doc, x)

    # ============ 4. UML 与数据建模 ============
    add_h1(doc, "4  UML 与数据建模")
    add_body(doc, "本章承载需求阶段的可视化建模产出，配套对应图表。")

    add_h2(doc, "4.1  系统用例图")
    add_body(doc,
        "主参与者：学生（STU）、教师（TCH）、自学者（DEV）。核心用例包括：浏览算法目录、"
        "选择算法、输入数据、随机生成数据、载入预置用例、播放/暂停/单步、查看当前步骤说明、"
        "查看复杂度信息、切换算法。其中“首次启动引导”作为 include 关系附属于“选择算法”。"
    )
    if os.path.exists(r"D:/课设/算法可视化/figures/fig_usecase.png"):
        doc.add_picture(r"D:/课设/算法可视化/figures/fig_usecase.png",
                        width=Cm(14.0))
        add_caption(doc, "图 4-1  系统用例图")

    add_h2(doc, "4.2  整体架构图 / 包图")
    add_body(doc,
        "采用三层 + 旁支的架构：表现层（Vue 单文件组件：Sidebar / Topbar / Canvas / "
        "Controls / DataPanel / StepPanel / ComplexityPanel）→ 领域层"
        "（algorithms/：图算法、排序、线性表、栈队列、树结构等）→ 基础设施层"
        "（router/、store/、utils/、theme/）；状态管理（Pinia）作为旁支贯通三层。"
        "三层之间通过响应式数据流与事件回调通信，避免单向耦合，保证模块可独立替换。"
    )

    add_h2(doc, "4.3  核心业务类图")
    add_body(doc,
        "核心类：Algorithm（抽象基类，含 name、difficulty、timeComplexity、spaceComplexity、"
        "code、init、step、reset）；GraphAlgorithm（Dijkstra 等）、SortingAlgorithm"
        "（冒泡 / 快速）、LinearListAlgorithm、TreeAlgorithm、StackQueueAlgorithm 等具体子类；"
        "Tracer（步骤记录器）；Stepper（播放状态机：idle/running/paused/completed）；"
        "TestCase（输入数据 + 元信息）。"
        "类间关系：Algorithm 与 Tracer 是聚合关系；"
        "Stepper 与 Algorithm 是关联关系（一个 Stepper 可驱动一个 Algorithm 的步骤序列）。"
    )

    add_h2(doc, "4.4  关键流程活动图")
    add_body(doc,
        "下图给出系统最核心的“算法播放主流程”活动图，覆盖从启动应用、选择算法、数据校验、"
        "构建步骤序列、渲染画布、按用户操作分支推进帧、直到全部步骤完成的全过程。"
        "决策节点（菱形）出现在“是否随机数据 / 校验通过 / 操作类型 / 是否末步”四处，"
        "充分体现了播放控制的状态分支。"
    )
    if os.path.exists(r"D:/课设/算法可视化/figures/fig_activity.png"):
        doc.add_picture(r"D:/课设/算法可视化/figures/fig_activity.png",
                        width=Cm(12.0))
        add_caption(doc, "图 4-2  算法播放主流程活动图")

    add_h2(doc, "4.5  数据流图（DFD）")
    add_body(doc,
        "为表达系统内外的数据搬运过程，采用 Yourdon-DeMarco 风格的数据流图（DFD）。"
        "外部实体包括“用户”“浏览器屏幕”“浏览器存储”“导出数据”；"
        "内部处理 P1–P5 涵盖数据输入与校验、算法执行与追踪、渲染与同步、用例与日志管理、导出引擎；"
        "数据存储 D1（用例库 TestCase）和 D2（执行日志 ExecutionLog）。"
        "数据流以箭头命名表达流向与含义，整体形成完整闭环，无悬空数据流。"
    )
    if os.path.exists(r"D:/课设/算法可视化/figures/fig_dfd.png"):
        doc.add_picture(r"D:/课设/算法可视化/figures/fig_dfd.png",
                        width=Cm(14.0))
        add_caption(doc, "图 4-3  数据流图（DFD 一层展开）")

    add_h2(doc, "4.6  状态机图（播放控制）")
    add_body(doc,
        "状态：idle、running、paused、completed。事件与转换："
        "idle —play→ running； idle —stepForward→ paused； running —pause→ paused； "
        "paused —play→ running； running —last step done→ completed； "
        "completed —reset→ idle； 任意 —reset→ idle。"
    )
    if os.path.exists(r"D:/课设/算法可视化/figures/fig_state.png"):
        doc.add_picture(r"D:/课设/算法可视化/figures/fig_state.png",
                        width=Cm(13.0))
        add_caption(doc, "图 4-4  播放控制状态机图")

    add_h2(doc, "4.7  数据模型设计：E-R 图与 DBML")
    add_body(doc,
        "本系统的本地持久化数据建模为 4 个实体：Algorithm（算法元数据，运行时常量）、"
        "TestCase（用户保存的输入用例）、ExecutionLog（每次运行产生的执行日志）、"
        "Preference（用户偏好键值对）。其中 Algorithm ↔ TestCase 为一对多“拥有”关系；"
        "Algorithm ↔ ExecutionLog 为一对多“产生”关系；TestCase ↔ ExecutionLog 为一对多“包含”关系。"
    )
    if os.path.exists(r"D:/课设/算法可视化/figures/fig_er.png"):
        doc.add_picture(r"D:/课设/算法可视化/figures/fig_er.png",
                        width=Cm(14.0))
        add_caption(doc, "图 4-5  本地数据 E-R 图")

    add_body(doc, "若启用本地持久化（如 IndexedDB），主要表的 DBML 描述如下：")
    dbml = (
        "Table TestCase {\n"
        "  id           int [pk, increment]\n"
        "  algorithm_id varchar [not null]   // 算法标识，如 'graph.dijkstra'\n"
        "  name         varchar [not null]   // 用例名称\n"
        "  payload      text    [not null]   // 输入数据（JSON 字符串）\n"
        "  created_at   datetime [default: `now()`]\n"
        "  updated_at   datetime\n"
        "  Indexes { (algorithm_id, name) [unique] }\n"
        "}\n\n"
        "Table ExecutionLog {\n"
        "  id           int [pk, increment]\n"
        "  algorithm_id varchar [not null]\n"
        "  case_id      int [ref: > TestCase.id, null]\n"
        "  steps        text [not null]      // 步骤序列（JSON）\n"
        "  duration_ms  int\n"
        "  created_at   datetime [default: `now()`]\n"
        "}\n\n"
        "Table Preference {\n"
        "  key   varchar [pk]   // theme / animation / shortcut / language ...\n"
        "  value text\n"
        "}"
    )
    add_code_block(doc, dbml)

    # ============ 5. AI 辅助 ============
    add_h1(doc, "5  AI 辅助需求分析的应用")
    add_body(doc,
        "本项目在需求阶段以 AI 工具作为协同生产力，把 AI 用作“能 7×24 小时陪同思考的需求工程师”，"
        "覆盖需求挖掘、用例生成、风险扫描、UML/DBML 建模、文档润色等环节。"
        "团队约定：AI 仅生成草稿与候选答案，所有产出必须经人工评审、复核与定稿后才会进入文档。"
    )

    add_h2(doc, "5.1  应用目标与原则")
    obj = [
        "原则一“放大不替代”：AI 用于放大团队的需求覆盖度与产出速度，但不替代团队对业务的判断与决策。",
        "原则二“可追溯”：所有 AI 产出都登记在文档版本记录中，并附人工复核记录，便于回溯。",
        "原则三“分阶段使用”：需求挖掘 → 用例生成 → 风险扫描 → 建模辅助 → 文案润色，分阶段调用，避免“一锅烩”。",
        "原则四“数据脱敏”：禁止把团队成员个人信息、未公开课题数据上传给在线 AI；本地知识脱敏后再使用。",
    ]
    for o in obj:
        add_bullet(doc, o)

    add_h2(doc, "5.2  应用场景与对应工具")
    tool_rows = [
        ["需求拆解", "把模糊的“做一个算法可视化”愿景拆解为 7 个功能模块、30+ 条 FR。",
         "Trae IDE / ChatGPT", "采用 60% / 修订 30% / 剔除 10%"],
        ["用例与边界条件生成", "针对每个算法批量生成正常/边界/异常输入。",
         "Claude Code", "采用 65%"],
        ["风险点识别", "枚举浏览器兼容、动画性能、字体渲染、依赖供应链等可能踩坑点。",
         "ChatGPT", "采用 50%"],
        ["UML / DBML 建模", "依据自然语言需求生成用例图、类图、E-R 图骨架与 DBML 表。",
         "Claude Code + python-docx", "采用 70%"],
        ["文档润色", "对功能描述、错误提示、引导文案统一风格、纠正语病。",
         "ChatGPT", "采用 80%"],
        ["可行性论证", "从经济、操作、合规多维度补全可行性研究的论证逻辑。",
         "ChatGPT", "采用 55%"],
    ]
    add_table(doc,
              ["应用场景", "AI 在这里做什么", "使用工具", "采纳比例"],
              tool_rows, col_widths=[3.0, 6.0, 3.5, 3.0])

    add_h2(doc, "5.3  工作流与示范提示词（Prompt）")
    add_body(doc,
        "团队对 AI 的使用统一遵循“四步工作流”：① 设定角色与上下文 → ② 给出明确目标与产出格式 → "
        "③ 提供输入材料与约束 → ④ 要求自我反思与校对。下面给出几条实际使用过的提示词模板。"
    )

    add_h3(doc, "5.3.1  提示词 P1：需求拆解")
    add_code_block(doc,
        "你是一名资深需求分析师。请基于以下产品愿景：\n"
        "「面向《数据结构》课程的 Web 端算法可视化实验平台，覆盖排序/线性表/栈队列/树/图/递归等算法，\n"
        "提供专用画布、源码级步骤同步与测试数据闭环」。\n"
        "请输出：① 不少于 6 个一级功能模块；② 每个模块下 3~5 条 IEEE 830 风格的 FR-XXX；\n"
        "③ 用 Markdown 表格输出；④ 用结尾段说明你不确定的需求点（不少于 3 条）。"
    )

    add_h3(doc, "5.3.2  提示词 P2：用例与边界条件生成")
    add_code_block(doc,
        "针对“Dijkstra 最短路径可视化”模块，请生成 8 条测试用例，覆盖：\n"
        "① 常规连通图（5/6/4 节点）；② 极端图（n=1、n=2、孤立节点、起点等于终点）；\n"
        "③ 异常输入（含字母权重、空输入、起点超出节点数）。\n"
        "对每条用例输出：用例名 / 输入数据 / 预期可视化行为 / 预期最短路径 / 说明。"
    )

    add_h3(doc, "5.3.3  提示词 P3：风险扫描")
    add_code_block(doc,
        "你是一名 Vue.js + Vite 资深前端工程师。请列出本项目可能遇到的 10 个高风险点（含浏览器兼容、\n"
        "Canvas 性能、依赖供应链、动画卡顿、移动端适配等），对每个风险输出：\n"
        "现象 / 触发条件 / 严重程度（高/中/低）/ 缓解方案 / 备选方案。"
    )

    add_h3(doc, "5.3.4  提示词 P4：UML / DBML 骨架生成")
    add_code_block(doc,
        "请基于以下需求摘要：「平台需要本地保存测试用例、执行日志、用户偏好；用例归属算法；\n"
        "执行日志可关联用例，亦可独立存在；偏好使用键值对存储」，输出 DBML 表结构（含主外键、索引），\n"
        "并附 PlantUML 用例图脚本，描述学生/教师/自学者对系统的核心用例。"
    )

    add_h2(doc, "5.4  AI 辅助产出清单（实物）")
    out = [
        "FR 编号清单初稿：AI 给出 42 条，人工修改 + 扩展至 30+ 条 FR。",
        "Algorithm 抽象基类 JavaScript 接口骨架：含 8 个公开方法签名与文档注释。",
        "测试用例边界条件清单：覆盖 6 类算法、共计 60+ 条用例草案，最终选 30+ 条精修后入库。",
        "兼容性问题预案：浏览器 Canvas 性能差异、字体渲染、依赖版本冲突等 12 条预案。",
        "本文档第 2/3/4/6 章的部分二级标题与文案初稿。",
    ]
    for o in out:
        add_bullet(doc, o)

    add_h2(doc, "5.5  效果验证与人工复核机制")
    verify = [
        "人工复核分工：组长（孟欣）逐条核对需求与功能边界、剔除 AI 幻觉项；"
        "李佳乐复核 UI/交互可行性与文案风格；宋启悦复核测试用例覆盖度与边界条件；"
        "冀文卓复核技术选型、版本号、构建链路与法律合规。",
        "采纳率统计：AI 草稿被原样采纳约 60%，需要修改后采纳约 30%，被剔除约 10%（多为虚构 API、"
        "不准确版本号、与课程边界冲突的功能）。",
        "三级评审：AI 产出 → 单人复核 → 组内交叉评审 → 组长定稿，确保产出质量。",
        "结论：AI 辅助显著提升了需求梳理的效率与覆盖度，但不能替代领域判断与团队对齐；"
        "在课程设计这种高强度短周期项目中，AI 的“扩展思路、补全细节、统一文风”三种能力价值最高。",
    ]
    for v in verify:
        add_bullet(doc, v)

    add_h2(doc, "5.6  风险与边界")
    risk = [
        "幻觉风险：AI 可能给出虚构的 API、版本号或文档链接，需通过官方文档二次核实。",
        "版权风险：AI 生成的代码片段可能与开源代码相似，需用 GitHub 搜索或代码相似度工具自查。",
        "上下文截断风险：长文档容易超出模型上下文，团队约定单次提问产出不超过 2000 字，分批迭代。",
        "数据合规风险：禁止将含个人信息、内部文档的内容直接发送给在线 AI，必须先脱敏。",
    ]
    for r in risk:
        add_bullet(doc, r)

    # ============ 6. 附录 ============
    add_h1(doc, "6  附录")

    add_h2(doc, "6.1  算法全量清单")
    cat_rows = [
        ["线性表", "顺序表-数组（简单）；单链表-不带头节点（简单）；单链表-带头结点（简单）；"
                  "双链表-不带头结点（中等）；双链表-带头结点（中等）；循环单链表（中等）；循环双链表（中等）"],
        ["栈和队列", "栈-顺序表（简单）；栈-链表（简单）；队列-顺序表（简单）；"
                    "队列-链表（简单）；循环队列（中等）；括号匹配（简单）；表达式计算（中等）"],
        ["树结构", "二叉树-链式存储（简单）；二叉排序树（中等）；哈夫曼树（中等）；线索二叉树（中等）"],
        ["图结构", "邻接矩阵（简单）；邻接链表（简单）；BFS（中等）；DFS（中等）；"
                  "Dijkstra 最短路径（中等+）"],
        ["排序", "冒泡排序（简单）；快速排序（中等）；直接插入排序（简单）；"
                "希尔排序（中等）；简单选择排序（简单）；归并排序（中等）"],
        ["递归与回溯", "汉诺塔（中等）；棋盘覆盖（中等）；图着色（中等）；迷宫求解（中等）"],
    ]
    add_table(doc, ["分类", "算法清单（含难度）"], cat_rows,
              col_widths=[2.5, 13.0])

    add_h2(doc, "6.2  术语表")
    term_rows = [
        ["专用画布", "针对每一类数据结构使用对应渲染策略的可视化区域，不套用统一模板。"],
        ["源码级步骤同步", "每一帧动画与一行真实源码、一条自然语言步骤说明三方绑定。"],
        ["测试数据闭环", "支持随机生成、手动输入、用例载入、结果观察的完整数据通路。"],
        ["难度徽章", "标注算法对学习者的认知难度，分简单 / 中等 / 中等+ / 高 四档。"],
        ["播放速度", "单步切换之间的等待毫秒数，可在播放控制面板的滑块中调节。"],
        ["状态机", "由 idle / running / paused / completed 四个状态构成的播放控制模型。"],
    ]
    add_table(doc, ["术语", "释义"], term_rows, col_widths=[3.5, 12.0])

    add_h2(doc, "6.3  参考文献与引用规范")
    refs = [
        "[1] OMG, Unified Modeling Language Specification, Version 2.5.1, 2017.",
        "[2] DBML, Database Markup Language Specification 1.0, https://dbml.dbdiagram.io/.",
        "[3] IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications.",
        "[4] GB/T 8567-2006, 计算机软件文档编制规范.",
        "[5] 太原理工大学软件学院. 软件工程专业毕业设计（论文）写作规范与要求(2026).",
        "[6] Cormen T H, Leiserson C E, Rivest R L, et al. 算法导论（第 3 版）[M]. 北京: 机械工业出版社, 2013.",
        "[7] 严蔚敏, 吴伟民. 数据结构（C 语言版）[M]. 北京: 清华大学出版社, 2011.",
        "[8] Vue.js Documentation, https://cn.vuejs.org/.",
        "[9] Vite Documentation, https://cn.vitejs.dev/.",
        "[10] Pinia Documentation, https://pinia.vuejs.org/zh/.",
        "[11] MDN Web Docs, Canvas API. https://developer.mozilla.org/zh-CN/docs/Web/API/Canvas_API.",
        "[12] WCAG 2.1, Web Content Accessibility Guidelines, W3C, 2018.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(r)
        set_run_font(run, "宋体", "Times New Roman", size=SZ_WU_HAO)

    add_h2(doc, "6.4  文档版本记录")
    ver_rows = [
        ["v0.1", "2026-06-19", "孟欣",
         "初稿：完成第 1~6 章基础内容，建立全部 FR 编号与系统目标 G 编号。"],
        ["v0.2", "2026-06-19", "李佳乐",
         "补充非功能需求（性能/兼容/无障碍）与可行性研究 2.2 / 2.3 节。"],
        ["v0.3", "2026-06-19", "宋启悦",
         "完善 UML 与数据建模章节，补充用例边界与术语表。"],
        ["v0.4", "2026-06-19", "冀文卓",
         "补充技术可行性表格、AI 辅助章节与参考文献。"],
        ["v1.0", "2026-06-19", "孟欣",
         "全文按学校论文规范统稿，调整字体字号字重与页边距，提交团队评审。"],
    ]
    add_table(doc, ["版本", "日期", "作者", "变更说明"],
              ver_rows, col_widths=[1.6, 2.6, 2.5, 8.8])

    # 团队签名
    for _ in range(2):
        doc.add_paragraph()
    add_h2(doc, "团队签名确认")
    sig_rows = [
        ["孟欣（组长）", "________", "________"],
        ["李佳乐", "________", "________"],
        ["宋启悦", "________", "________"],
        ["冀文卓", "________", "________"],
    ]
    add_table(doc, ["成员", "签名", "日期"],
              sig_rows, col_widths=[5.0, 5.0, 5.0])

    out_path = r"D:/课设/算法可视化/需求分析文档.docx"
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    main()
