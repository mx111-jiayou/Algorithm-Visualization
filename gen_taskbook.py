# -*- coding: utf-8 -*-
"""
生成《算法过程可视化系统》课程设计详细任务书
基于 D:/课设/算法可视化/要求/任务/ 下 1.md~15.md 的内容整合
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')


def add_para(doc, text, font="宋体", size=12, bold=False, align=None,
             first_line_indent=True, space_before=0, space_after=0,
             line_spacing=1.5, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, font, size, bold, color)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, "黑体", 15, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "黑体", 14, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, "黑体", 12, bold=True)
    return p


def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("● " + text)
    set_run_font(run, "宋体", size)
    return p


def main():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(3.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.3)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ============ 封面 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("《程序设计课程设计》")
    set_run_font(run, "黑体", 26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("课设任务书")
    set_run_font(run, "黑体", 22, bold=True)

    # 留白
    for _ in range(2):
        doc.add_paragraph()

    info_lines = [
        ("题    目：", "算法过程可视化系统"),
        ("班    级：", "软件 2301"),
        ("组    长：", "_____________________"),
        ("组    员：", "_____________________"),
        ("指导老师：", "_____________________"),
        ("起止时间：", "_____________________"),
    ]
    for label, val in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(label + val)
        set_run_font(run, "宋体", 16, bold=False)

    # 留白
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("软件学院")
    set_run_font(run, "宋体", 16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("二〇二六年")
    set_run_font(run, "宋体", 16, bold=True)

    doc.add_page_break()

    # ============ 任务书正文 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("任 务 书")
    set_run_font(run, "黑体", 22, bold=True)
    p.paragraph_format.space_after = Pt(18)

    # ====== 一、课题背景与意义 ======
    add_heading1(doc, "一、课题背景与意义")

    add_heading2(doc, "1.1 课题背景")
    add_para(doc,
        "算法是计算机科学的核心基石，但传统教学中算法常以伪代码、流程图或抽象文字描述呈现，"
        "学习者难以在脑海中构建动态执行过程，从而陷入“看得懂语法、说不清原理”的困境。"
        "排序算法的元素交换、最短路径算法的松弛过程、递归算法的栈式调用，这些动态行为如能借助"
        "图形动画直观展示，将极大降低理解门槛、提升学习兴趣。"
    )
    add_para(doc,
        "随着 Web 前端技术（HTML5 + Canvas + JavaScript）以及 Python 桌面 GUI（Tkinter、PyQt）"
        "的成熟，构建跨平台、可交互的算法可视化系统在技术上已完全可行。本课题正是在这一背景下提出，"
        "要求设计并实现一个集排序算法、图算法、递归算法于一体的过程可视化系统，"
        "通过统一的界面框架与执行控制器，让用户能够分步、动态地观察典型算法的内部运行过程。"
    )

    add_heading2(doc, "1.2 课题目的")
    add_bullet(doc, "综合训练学生数据结构、算法、面向对象、GUI 程序设计与软件工程的整体能力。")
    add_bullet(doc, "深入理解快速排序、Dijkstra 最短路径、汉诺塔等经典算法的核心思想与执行细节。")
    add_bullet(doc, "掌握模块化设计、状态机驱动、观察者模式、主题统一等工程化实践方法。")
    add_bullet(doc, "完成一个具备实用价值、可对外演示的算法可视化教学辅助系统。")

    add_heading2(doc, "1.3 课题意义")
    add_para(doc,
        "本课题既是一次算法学习的“再深化”，也是一次软件项目从需求分析、概要设计、详细设计、"
        "编码实现到系统测试与文档编写的完整工程演练。完成本课题后，学生应当具备独立设计中等规模"
        "桌面/Web 应用程序的能力，并能熟练运用动画、交互手段对抽象算法进行直观表达。"
    )

    # ====== 二、设计任务与功能要求 ======
    add_heading1(doc, "二、设计任务与总体功能要求")

    add_heading2(doc, "2.1 设计任务总体描述")
    add_para(doc,
        "本课题要求设计并实现《算法过程可视化系统》。系统应具备统一的主界面，"
        "支持用户从菜单中选择不同算法（排序、图、递归），输入或随机生成测试数据，"
        "执行算法并以图形动画的方式展示算法的关键步骤、当前状态以及最终结果，"
        "同时提供单步、自动播放、暂停、重置等执行控制功能，并显示对应的时间/空间复杂度信息。"
    )

    add_heading2(doc, "2.2 总体功能模块划分")
    add_bullet(doc, "（1）系统架构与统一界面框架：技术选型、环境搭建、主界面设计、数据输入与随机生成。")
    add_bullet(doc, "（2）排序算法可视化模块：快速排序核心逻辑、过程动画、最终结果与复杂度输出。")
    add_bullet(doc, "（3）图算法可视化模块：图数据结构与输入、Dijkstra 算法核心逻辑、图形化路径展示。")
    add_bullet(doc, "（4）递归与回溯算法可视化模块：汉诺塔核心递归、过程动画、测试用例与复杂度说明。")
    add_bullet(doc, "（5）系统集成与扩展功能：算法执行控制、模块集成与界面优化、系统测试与文档编写。")

    add_heading2(doc, "2.3 总体性能与质量要求")
    add_bullet(doc, "可视化动画流畅，无明显闪烁、卡顿；典型规模下单步切换响应时间不超过 200 ms。")
    add_bullet(doc, "界面布局合理、风格统一，配色方案、字体字号在各模块之间保持一致。")
    add_bullet(doc, "代码结构清晰，遵循模块化、面向对象设计原则，关键类/函数具有必要注释。")
    add_bullet(doc, "对非法输入、空数据、边界数据具备完整的容错处理与友好的错误提示。")
    add_bullet(doc, "项目根目录提供 README 文档，第三方在干净环境下可按文档复现运行结果。")

    # ====== 三、各章节详细任务要求 ======
    add_heading1(doc, "三、各章节详细任务要求")

    # 第 1 章
    add_heading2(doc, "3.1 第 1 章 系统架构设计与统一界面框架")

    add_heading3(doc, "3.1.1 任务 1：技术选型与环境搭建")
    add_bullet(doc, "对比 Web 前端方案（HTML/CSS/JS + Canvas）与 Python 桌面方案（Tkinter/PyQt）的优缺点。")
    add_bullet(doc, "完成技术选型论证，明确所选方案的理由（视觉效果、部署成本、交互能力等维度）。")
    add_bullet(doc, "完成开发环境搭建：浏览器（Chrome/Firefox）+ VS Code + Live Server 插件，或 Python 3.9+ 解释器与依赖库。")
    add_bullet(doc, "完成第一个 Hello World 程序：在 Canvas 或主窗口上绘制矩形、圆形、文本等基本图元，验证环境可用。")

    add_heading3(doc, "3.1.2 任务 2：统一界面设计与入口实现")
    add_bullet(doc, "采用模块化布局思想，将主界面划分为四个功能区：算法选择区、数据输入区、过程展示区、结果信息面板。")
    add_bullet(doc, "实现算法选择下拉菜单（Combobox），支持“快速排序”“Dijkstra”“汉诺塔”等选项的切换。")
    add_bullet(doc, "实现数据输入文本框与“随机生成”按钮；实现可视化画布（Canvas）以及步骤/复杂度信息面板。")
    add_bullet(doc, "保证一致性、可维护性与可扩展性：所有算法共享同一套控件与执行流程“选择→输入→运行→展示”。")

    add_heading3(doc, "3.1.3 任务 3：数据输入与随机生成模块")
    add_bullet(doc, "设计统一数据输入接口：数组采用 [a, b, c] 形式，图采用邻接矩阵 [[…],[…]] 形式。")
    add_bullet(doc, "实现 parse_array_input、parse_graph_input 解析函数，将字符串转换为可计算的数据结构。")
    add_bullet(doc, "实现 validate_input 验证函数，能识别括号不匹配、非法字符、矩阵非方阵、权重越界等错误并给出友好提示。")
    add_bullet(doc, "实现 generate_random_array、generate_random_graph 随机生成器，支持长度、范围、密度、权重区间等参数。")

    # 第 2 章
    add_heading2(doc, "3.2 第 2 章 排序算法可视化实现")

    add_heading3(doc, "3.2.1 任务 4：快速排序算法核心逻辑实现")
    add_bullet(doc, "理解并阐述快速排序的分治思想（选基准 → 分区 → 递归）。")
    add_bullet(doc, "采用 Lomuto 分区方案实现 partition(arr, low, high) 函数，正确返回基准最终索引。")
    add_bullet(doc, "实现递归 quick_sort(arr, low, high) 函数，对左右子数组分别递归。")
    add_bullet(doc, "在算法执行过程中，将每一次比较、交换、基准归位事件按 {type, indices, values} 结构追加到 steps 列表，为可视化提供数据。")

    add_heading3(doc, "3.2.2 任务 5：排序过程动画与步骤展示")
    add_bullet(doc, "实现 StepCollector 类（或等价模块），在比较、交换、指针移动等关键节点保存数组快照、左右指针、基准索引、说明文本等。")
    add_bullet(doc, "在 Canvas 上绘制柱状图：柱宽 = 画布宽 / 数组长度，柱高与元素值成比例；不同状态（普通、比较中、交换中、基准）以不同颜色高亮。")
    add_bullet(doc, "在画布上绘制左指针、右指针、基准（PIVOT）等标记，并提供步骤计数器与文字说明区。")
    add_bullet(doc, "提供“上一步 / 下一步 / 自动播放 / 速度滑块”等交互控件，实现分步与连续两种播放模式。")

    add_heading3(doc, "3.2.3 任务 6：最终结果与复杂度信息输出")
    add_bullet(doc, "在排序流程中加入 stats（comparisons、swaps）计数器，每轮算法启动前重置。")
    add_bullet(doc, "实现 displayResult、displayStats、displayComplexity 三个展示函数，分别输出排序结果数组、关键步骤统计、时间/空间复杂度文字说明。")
    add_bullet(doc, "在界面信息面板中清晰展示：时间复杂度（平均 O(n log n)，最坏 O(n²)）、空间复杂度（O(log n)）。")

    # 第 3 章
    add_heading2(doc, "3.3 第 3 章 图算法可视化实现")

    add_heading3(doc, "3.3.1 任务 7：图数据结构与输入模块")
    add_bullet(doc, "设计统一的 Graph 抽象类，定义 add_edge、get_weight、get_neighbors 等接口。")
    add_bullet(doc, "实现 AdjacencyMatrixGraph 与 AdjacencyListGraph 两种存储结构，并能阐述其空间复杂度与适用场景。")
    add_bullet(doc, "实现手动输入功能：用户可输入节点数、边数及每条边的 (u, v, weight)。")
    add_bullet(doc, "实现随机生成功能：先生成生成树以保证连通性，再随机添加额外边并赋权重，避免自环和重复边。")

    add_heading3(doc, "3.3.2 任务 8：Dijkstra 算法核心逻辑实现")
    add_bullet(doc, "理解 Dijkstra 算法的贪心思想与“松弛”操作；说明其在非负权图上的正确性。")
    add_bullet(doc, "采用 heapq 优先队列实现 dijkstra(graph, source) 函数，输出 dist、prev 字典。")
    add_bullet(doc, "在每一轮“出队 + 标记 visited + 松弛邻居”操作后，向 steps 列表追加快照（current_node, dist, visited, edge_relaxed）。")
    add_bullet(doc, "提供路径回溯函数：根据 prev 从目标节点反推回源节点，得到最短路径序列。")

    add_heading3(doc, "3.3.3 任务 9：图可视化与路径展示")
    add_bullet(doc, "使用 matplotlib + networkx（或 Canvas 自绘）绘制图结构：节点以圆形表示，边以带权重标签的线段表示。")
    add_bullet(doc, "采用 spring_layout 或固定布局确定节点坐标，避免绘图重叠。")
    add_bullet(doc, "实现 update_graph_visualization：根据 steps 中的 current_node、visited、edge_relaxed 动态修改节点和边的颜色、粗细。")
    add_bullet(doc, "算法结束后，将最终最短路径以独立颜色（如红色加粗）高亮显示，并在界面上标注路径总长度。")

    # 第 4 章
    add_heading2(doc, "3.4 第 4 章 递归与回溯算法可视化")

    add_heading3(doc, "3.4.1 任务 10：汉诺塔算法核心逻辑实现")
    add_bullet(doc, "用自己的语言阐述递归基线条件与递归步骤；理解汉诺塔“移动 n-1 → 移动最大盘 → 移动 n-1”的拆解。")
    add_bullet(doc, "实现 hanoi(n, source, target, auxiliary, steps) 函数，正确产生 2ⁿ-1 步移动序列。")
    add_bullet(doc, "约定盘子编号规则（1 号最小、n 号最大），将每一步以 (盘子编号, 源柱, 目标柱) 元组追加至 steps。")

    add_heading3(doc, "3.4.2 任务 11：汉诺塔过程可视化展示")
    add_bullet(doc, "设计柱子状态数据结构 towers = {'A': […], 'B': […], 'C': […]}，列表末尾代表柱顶。")
    add_bullet(doc, "实现 draw_towers 函数：在画布上绘制三根柱子和相应大小的盘子（建议使用矩形/圆角矩形，宽度与盘号成正比，颜色区分）。")
    add_bullet(doc, "在递归 hanoi 函数中插入“暂停 + 刷新”逻辑（time.sleep + 重绘 / FuncAnimation），形成可视动画。")
    add_bullet(doc, "实现步骤列表组件，实时展示每步移动指令，并显示总移动次数。")

    add_heading3(doc, "3.4.3 任务 12：测试用例与复杂度说明")
    add_bullet(doc, "为三类算法设计共计不少于 6 条测试用例（每类至少 2 条），覆盖小规模、中规模、边界（空/单元素/起点=终点/1 个盘子）等情况。")
    add_bullet(doc, "在主界面增加“测试用例选择器”，选择后自动填充输入框并可一键执行。")
    add_bullet(doc, "在界面预留复杂度展示区，根据当前算法显示固定的时间/空间复杂度，并显示当前规模与实际执行步数：")
    add_bullet(doc, "    排序：O(n log n) / O(log n)；Dijkstra：O((V+E) log V) / O(V)；汉诺塔：O(2ⁿ) / O(n)。")

    # 第 5 章
    add_heading2(doc, "3.5 第 5 章 系统集成与扩展功能")

    add_heading3(doc, "3.5.1 任务 13：算法执行控制功能实现")
    add_bullet(doc, "采用状态机思想设计 AlgorithmController，包含 idle / running / paused / completed 四种状态。")
    add_bullet(doc, "实现 stepForward、play、pause、reset 四个控制方法，并明确各状态下的合法转换。")
    add_bullet(doc, "通过 onStep、onReset 回调与可视化视图通信，实现控制器与视图的解耦。")
    add_bullet(doc, "在界面上根据当前状态自动启用/禁用按钮（如已完成时禁用单步与播放）。")

    add_heading3(doc, "3.5.2 任务 14：系统集成与界面优化")
    add_bullet(doc, "实现 MainController（或主窗口类），统一管理三个算法可视化模块的生命周期。")
    add_bullet(doc, "切换算法时：销毁旧模块、清空主框架子控件、实例化并 pack 新模块、重置状态。")
    add_bullet(doc, "建立 theme.py（或等价模块）统一管理主色、辅色、字体、字号、间距等主题信息。")
    add_bullet(doc, "运行算法时禁用算法切换下拉框；提供顶部状态栏显示“当前算法 + 当前状态”；为常用操作绑定键盘快捷键（可选）。")

    add_heading3(doc, "3.5.3 任务 15：系统测试与 README 文档编写")
    add_bullet(doc, "按排序、图、递归、执行控制四个维度，逐项执行功能测试与边界/异常测试，并形成测试记录。")
    add_bullet(doc, "在项目根目录编写 README.md，包含：项目简介、功能列表、技术栈、环境要求、安装与运行步骤、输入格式说明、项目结构、许可证等。")
    add_bullet(doc, "确保 README 中的安装与运行步骤可在干净的机器上完整复现；移除调试用打印语句与无效注释。")

    # ====== 四、输入输出与接口要求 ======
    add_heading1(doc, "四、输入输出与接口要求")

    add_heading2(doc, "4.1 输入格式")
    add_bullet(doc, "排序算法：以方括号包裹、逗号分隔的整数序列，如 [5, 2, 8, 1, 9]，仅支持整数。")
    add_bullet(doc, "Dijkstra 算法：邻接矩阵（如 [[0,2,4],[2,0,1],[4,1,0]]）或 “u v weight” 边列表两种形式之一，并指定起点与终点。")
    add_bullet(doc, "汉诺塔算法：单个正整数 n 表示盘子数量，建议 1 ≤ n ≤ 10。")

    add_heading2(doc, "4.2 输出要求")
    add_bullet(doc, "可视化区域：以动画方式展示算法关键步骤；用颜色/线宽/标记区分不同状态。")
    add_bullet(doc, "步骤面板：以文字描述当前步骤（如“比较索引 j 的元素 X 和基准 Y”、“更新节点 B 的距离为 4”）。")
    add_bullet(doc, "结果面板：显示最终结果（有序数组、最短路径与距离、移动序列）、关键步骤计数、时间/空间复杂度。")

    add_heading2(doc, "4.3 异常与边界处理要求")
    add_bullet(doc, "对空输入、单元素输入、含非法字符的输入应给出明确错误提示，不得直接抛出未捕获异常。")
    add_bullet(doc, "Dijkstra 应能处理孤立节点、起点等于终点等情况；汉诺塔应能处理 n = 0 与 n = 1。")

    # ====== 五、技术规范 ======
    add_heading1(doc, "五、技术规范与开发约束")
    add_bullet(doc, "推荐技术栈：HTML5 + CSS + JavaScript + Canvas API（Web 方案），或 Python 3.9+ + Tkinter/PyQt + matplotlib（桌面方案）。")
    add_bullet(doc, "代码组织：按功能模块拆分文件/类，主控制器与算法模块、视图与控制器之间需保持松耦合。")
    add_bullet(doc, "命名规范：标识符语义清晰，类采用大驼峰、函数与变量采用蛇形或小驼峰，并保持全工程统一。")
    add_bullet(doc, "注释规范：核心类与关键函数应有简明文档；不得堆砌冗余注释。")
    add_bullet(doc, "版本管理：建议使用 Git 进行版本控制，提交信息清晰，能反映开发演进过程。")

    # ====== 六、进度安排 ======
    add_heading1(doc, "六、进度安排建议")

    sched = [
        ("第 1 周", "完成第 1 章（任务 1-3）：技术选型、环境搭建、主界面框架、数据输入与随机生成模块。"),
        ("第 2 周", "完成第 2 章（任务 4-6）：快速排序核心逻辑、过程动画与最终结果输出。"),
        ("第 3 周", "完成第 3 章（任务 7-9）：图数据结构、Dijkstra 算法实现与图可视化展示。"),
        ("第 4 周", "完成第 4 章（任务 10-12）：汉诺塔实现与可视化、测试用例与复杂度说明。"),
        ("第 5 周", "完成第 5 章（任务 13-15）：执行控制、系统集成与界面优化、系统测试与 README 文档；整理课设报告并准备答辩。"),
    ]
    table = doc.add_table(rows=1 + len(sched), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, t in enumerate(["时间", "工作内容"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(t)
        set_run_font(run, "黑体", 12, bold=True)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (a, b) in enumerate(sched, start=1):
        cells = table.rows[i].cells
        for j, t in enumerate([a, b]):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            run = p.add_run(t)
            set_run_font(run, "宋体", 12)
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ====== 七、考核与提交要求 ======
    add_heading1(doc, "七、考核与提交要求")

    add_heading2(doc, "7.1 提交物清单")
    add_bullet(doc, "（1）系统源代码：完整的工程目录，含所有源文件、资源文件与第三方依赖说明（requirements.txt 或 package.json）。")
    add_bullet(doc, "（2）可运行的程序：启动脚本或可执行文件，能够直接演示三类算法的可视化过程。")
    add_bullet(doc, "（3）课程设计报告：依据《课设报告模板》撰写，包含需求分析、概要设计、详细设计、系统实现、系统测试等章节。")
    add_bullet(doc, "（4）个人开发日志：记录每位组员每日工作内容与解决的关键问题。")
    add_bullet(doc, "（5）README.md：包含项目简介、功能列表、技术栈、环境要求、安装与运行步骤、输入格式说明等。")

    add_heading2(doc, "7.2 评分维度（建议权重，仅供参考）")
    add_bullet(doc, "功能完整性（30%）：是否覆盖三类算法、是否实现执行控制与统一界面。")
    add_bullet(doc, "可视化效果（20%）：动画流畅度、风格一致性、交互友好性。")
    add_bullet(doc, "代码质量（20%）：模块化设计、命名规范、注释合理性、容错处理。")
    add_bullet(doc, "测试与文档（15%）：测试用例覆盖度、README 与课设报告质量。")
    add_bullet(doc, "答辩与协作（15%）：演示效果、问题回答情况、组内分工的合理性。")

    add_heading2(doc, "7.3 验收检查清单")
    add_bullet(doc, "三类算法均可在统一界面下选择、输入、运行并完成可视化展示。")
    add_bullet(doc, "执行控制（单步、播放、暂停、重置）功能正常，且按钮状态与算法状态一致。")
    add_bullet(doc, "切换算法时旧模块的图形与状态完全清除，新模块以干净的初始状态显示。")
    add_bullet(doc, "至少 6 条测试用例（含边界用例）能正常运行；非法输入有明确提示，不会导致程序崩溃。")
    add_bullet(doc, "项目根目录的 README 在另一台干净机器上可按文档复现运行。")

    # ====== 八、参考资料 ======
    add_heading1(doc, "八、参考资料")
    refs = [
        "[1] Thomas H. Cormen 等. 算法导论（第 3 版）. 北京: 机械工业出版社, 2013.",
        "[2] Robert Sedgewick, Kevin Wayne. 算法（第 4 版）. 北京: 人民邮电出版社, 2012.",
        "[3] Mark Lutz. Python 学习手册（第 5 版）. 北京: 机械工业出版社, 2018.",
        "[4] David Flanagan. JavaScript 权威指南（第 7 版）. 北京: 机械工业出版社, 2021.",
        "[5] MDN Web Docs. Canvas API 官方文档. https://developer.mozilla.org/zh-CN/docs/Web/API/Canvas_API",
        "[6] Matplotlib 官方文档. https://matplotlib.org/stable/contents.html",
        "[7] NetworkX 官方文档. https://networkx.org/documentation/stable/",
        "[8] Python 官方文档（heapq、tkinter）. https://docs.python.org/zh-cn/3/",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        run = p.add_run(r)
        set_run_font(run, "宋体", 12)

    # ====== 签字栏 ======
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("指导教师签名：__________________     日期：______年______月______日")
    set_run_font(run, "宋体", 12)

    out = r"D:/课设/算法可视化/任务书.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
