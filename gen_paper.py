# -*- coding: utf-8 -*-
"""生成《算法过程可视化系统 课程设计报告》Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_ALIGN
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
BLACK = RGBColor(0, 0, 0)

# ---------- 页面设置 (A4, 上3.3 下2.3 左2.8 右2.3 cm) ----------
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(3.3)
sec.bottom_margin = Cm(2.3)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.3)

# ---------- 字体辅助 ----------
def set_font(run, cn='宋体', en='Times New Roman', size=12, bold=False, color=BLACK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = en
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)
    rfonts.set(qn('w:eastAsia'), cn)

# Normal 默认字体
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def heading(text, level):
    p = doc.add_heading('', level=level)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_font(run, cn='黑体', en='SimHei', size=15, bold=True)
        p.alignment = WD_ALIGN.CENTER
    elif level == 2:
        set_font(run, cn='黑体', en='SimHei', size=14, bold=True)
        p.alignment = WD_ALIGN.LEFT
    else:
        set_font(run, cn='黑体', en='SimHei', size=12, bold=True)
        p.alignment = WD_ALIGN.LEFT
    return p

def para(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, cn='宋体', en='Times New Roman', size=12)
    return p

def shade(paragraph, fill="F4F1EC"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def code(text, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(2)
        r = cp.add_run(caption)
        set_font(r, cn='宋体', en='Times New Roman', size=10.5, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    shade(p)
    lines = text.split('\n')
    run = p.add_run(lines[0])
    set_font(run, cn='宋体', en='Consolas', size=9, color=RGBColor(0x1A, 0x1A, 0x1A))
    for ln in lines[1:]:
        run.add_break()
        run.add_text(ln)
    return p

def field(paragraph, instr, placeholder=''):
    r = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = placeholder
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    r._r.append(b); r._r.append(it); r._r.append(sep); r._r.append(t); r._r.append(end)
    return r

def caption_fig(text):
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(6)
    r = cp.add_run(text)
    set_font(r, cn='宋体', en='Times New Roman', size=10.5)

# ====================== 封面 ======================
for _ in range(3):
    doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN.CENTER
r = t.add_run('算法过程可视化系统'); set_font(r, cn='黑体', en='SimHei', size=26, bold=True)
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN.CENTER
r = t2.add_run('课程设计报告'); set_font(r, cn='黑体', en='SimHei', size=22, bold=True)
for _ in range(4):
    doc.add_paragraph()
for label in ['项目名称：算法过程可视化系统',
              '技术栈：Vue 3 + Pinia + Vite + Canvas',
              '团队成员：孟欣　李佳乐　宋启悦　冀文卓',
              '完成日期：2026 年 6 月']:
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN.CENTER
    rr = pp.add_run(label); set_font(rr, cn='宋体', en='Times New Roman', size=14)
doc.add_page_break()

# ====================== 目录 ======================
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN.CENTER
tp.paragraph_format.space_after = Pt(12)
r = tp.add_run('目　录'); set_font(r, cn='黑体', en='SimHei', size=22, bold=True)
toc_p = doc.add_paragraph()
field(toc_p, 'TOC \\o "1-3" \\h \\z \\u', '在此处右键选择“更新域”可生成目录（含页码与点引线）。')
doc.add_page_break()

# 让 Word 打开时自动更新域（目录/页码）
settings = doc.settings.element
uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true')
settings.append(uf)

# ====================== 页脚页码 ======================
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN.CENTER
field(fp, 'PAGE', '1')

# ====================== 一、需求分析 ======================
heading('一、需求分析', 1)

heading('1.1 项目背景', 2)
para('数据结构与算法是计算机科学的核心基础课程，但其内容高度抽象，传统教学多依赖静态的代码文本与文字描述，'
     '学生难以在头脑中重建算法运行时的动态过程，导致排序的元素交换、图的遍历顺序、递归的调用与回溯等关键环节理解困难。')
para('为缓解这一教学痛点，本项目设计并实现了一个基于 Web 的交互式算法过程可视化系统。系统在浏览器中以动画方式逐步演示算法的执行过程，'
     '并与对应的 C 语言源代码实时联动高亮，配合步骤说明、音效反馈等手段，帮助学习者直观地建立对算法执行流程与核心原理的认识。')
para('系统共覆盖线性表、栈和队列、树结构、图结构、排序、递归与回溯六大类，合计 32 个经典算法与数据结构，'
     '可作为课堂演示工具与课后自主学习平台使用。')

heading('1.2 功能需求分析', 2)
para('（1）算法分类与选择：按六大类组织全部 32 个算法，用户可在顶部分类导航与左侧算法列表中快速定位并加载目标算法。')
para('（2）数据输入：支持数组、图、字符串、数值、迷宫等多种输入类型，提供手动输入、随机生成与预置测试用例三种方式。')
para('（3）动画演示：在中央画布中按步骤动态绘制算法状态变化，不同状态（比较、交换、基准、访问、已排序等）以不同颜色区分。')
para('（4）播放控制：提供播放、暂停、单步前进、单步回退、重置以及播放速度调节（100~2000ms），并支持空格/方向键/R 等键盘快捷键。')
para('（5）代码同步高亮：右侧代码面板展示算法源码，并实时高亮当前执行行，实现“动画—代码”一一对应。')
para('（6）步骤说明：步骤面板按序列出全部步骤及文字说明，可点击任意步骤跳转到对应状态。')
para('（7）音效反馈：基于 Web Audio API，根据当前步骤的操作类型播放相应音效，强化操作感知。')
para('（8）复杂度信息：展示每个算法的时间复杂度、空间复杂度与难度等级。')

heading('1.3 非功能需求分析', 2)
para('（1）性能需求：画布渲染需适配窗口尺寸并保持流畅，单步与播放操作响应及时，算法切换近乎即时完成。')
para('（2）易用性需求：界面采用左侧控制、中间画布、右侧代码/步骤的三栏布局，符合使用直觉，新用户无需培训即可上手。')
para('（3）兼容性需求：面向现代浏览器（支持 ES Module、Canvas 2D 与 Web Audio API），适配常见桌面分辨率。')
para('（4）可维护性需求：算法逻辑、状态管理与视图渲染分层解耦，新增算法只需补充步骤生成函数、注册元数据并扩展渲染分支。')

heading('1.4 可行性分析', 2)
para('（1）技术可行性：系统采用 Vue 3、Pinia、Vite、Canvas 2D 与 Web Audio API，均为成熟稳定的开源技术，社区生态完善，'
     '完全能够支撑动画渲染、状态管理与音效反馈等需求。')
para('（2）经济可行性：全部使用开源技术栈，无需任何授权费用；仅需普通 PC 即可开发与运行，硬件与维护成本低。')
para('（3）操作可行性：操作流程为“选择算法 → 输入数据 → 播放观察”，配合播放控制与快捷键，使用简单灵活。')
para('综合技术、经济与操作三方面分析，本系统具备完整的可行性，可进入开发与实现阶段。')

# ====================== 二、概要设计 ======================
heading('二、概要设计', 1)

heading('2.1 系统总体架构设计', 2)
para('系统采用前端单页应用架构，整体分为三层：视图层、状态层与算法逻辑层。')
para('视图层由根组件 App.vue 及其子组件（可视化画布 VisualizationCanvas、代码面板 CodePanel、步骤面板 StepsPanel）构成，'
     '负责界面展示与用户交互。状态层由 Pinia 全局 Store（algorithmStore）承担，集中管理当前算法、步骤序列、当前步骤索引、播放状态与速度等。'
     '算法逻辑层由 algorithms 目录下的若干纯函数模块组成，负责将算法执行过程转换为可供渲染的“步骤序列”。')
para('三层之间单向依赖：视图层从 Store 读取响应式状态并触发动作，Store 调用算法逻辑层生成步骤数据，'
     '算法逻辑层不依赖任何界面代码，保证了良好的可测试性与可扩展性。')
caption_fig('图 2-1　系统总体三层架构')

heading('2.2 系统功能模块划分', 2)
para('（1）算法注册与调度模块（registry.js / executor.js）：维护全部算法的元数据，并根据算法 ID 将输入数据分发到对应的步骤生成函数。')
para('（2）步骤生成模块（sorting.js、graph.js、recursion.js、dataStructures.js）：以纯函数形式实现各类算法，输出统一格式的步骤序列。')
para('（3）可视化渲染模块（VisualizationCanvas.vue）：依据步骤类型选择相应的绘制函数，在 Canvas 上渲染当前状态。')
para('（4）交互控制模块（algorithmStore.js 与左侧控制面板）：管理播放状态机、定时推进、单步与速度调节，并响应键盘快捷键。')
para('（5）代码与步骤展示模块（CodePanel.vue、StepsPanel.vue）：同步高亮当前代码行、展示步骤列表并支持跳转。')
para('（6）音效反馈模块（soundService.js）：基于 Web Audio API 实时合成音调，根据步骤类型给予听觉反馈。')

heading('2.3 系统运行流程设计', 2)
para('系统的核心运行流程如下：用户在左侧选择算法并输入数据后，App.vue 调用 executeAlgorithm 将数据交给对应算法模块，'
     '生成一个完整的步骤序列并存入 Store。Store 以 currentStep 索引指向当前步骤，stepInfo 计算属性派生出当前步骤对象。')
para('当用户点击播放时，Store 启动定时器按播放速度递增 currentStep；每推进一步，画布组件通过侦听重新渲染当前状态，'
     '代码面板根据步骤的 codeLine 高亮对应行，步骤面板滚动到当前步骤，音效模块播放相应反馈。到达末尾时播放状态切换为已完成并播放完成音效。')
caption_fig('图 2-2　系统运行流程')

# ====================== 三、详细设计 ======================
heading('三、详细设计', 1)

heading('3.1 系统整体设计说明', 2)
para('系统以“步骤序列（steps）”为核心数据结构贯穿各模块。每个算法模块均返回一个步骤数组，数组中的每个元素是一个描述某一时刻状态的对象，'
     '至少包含 description（步骤说明）与 codeLine（对应代码行号），并依据算法类型携带 array、highlights、nodeStates、towers、board、tree 等字段。')
para('视图层不关心算法如何计算，只负责把“某一步的状态对象”渲染出来；Store 不关心状态如何绘制，只负责在步骤序列上前后移动。'
     '这种以数据为中心的设计使三层彻底解耦。')

heading('3.2 算法可视化模块设计', 2)
para('可视化渲染由 VisualizationCanvas.vue 统一承担。组件在挂载时根据设备像素比（devicePixelRatio）对画布进行高清适配，并监听窗口尺寸变化重绘。'
     '渲染主函数 render 根据当前步骤的 type 字段（或算法 ID）分发到具体绘制函数，如排序柱状图、图、汉诺塔、棋盘、迷宫、栈、队列、链表、树等。')
para('为保证语义清晰，系统约定了一套状态配色：默认、比较、交换、基准、当前、已访问/已排序等分别对应不同颜色，'
     '使学习者能够通过颜色快速识别每个元素当前所处的状态。')

heading('3.3 交互与控制模块设计', 2)
para('交互控制集中在 Pinia Store 中，以一个简单的状态机管理播放过程，状态包括 idle（空闲）、running（播放中）、paused（暂停）、completed（已完成）。')
para('播放通过 setInterval 定时器按 playSpeed 推进步骤；单步前进/回退直接修改 currentStep 并触发对应音效；'
     '速度调节在播放中会重启定时器以即时生效；重置则停止定时器并将索引归零。App.vue 监听键盘事件，将空格、左右方向键、R 映射到相应动作。')

heading('3.4 数据输入与管理模块设计', 2)
para('系统按算法的 inputType 字段提供差异化的输入界面：array 接受逗号分隔的数字；graph 接受节点数、源点与每行一条的“起点,终点,权重”；'
     'string 接受括号串或表达式；number 接受盘子数或棋盘阶数；maze 直接随机生成迷宫。')
para('输入经 App.vue 中的解析函数（如 applyArrayInput、applyGraphInput）转换为标准数据结构后交给执行器。'
     '此外，generateRandomData 可按类型生成合理的随机数据，registry 中的 testCases 提供预置测试用例，形成完整的数据输入闭环。')

heading('3.5 界面设计', 2)
para('界面采用三栏式布局：顶部为分类导航栏，左侧为算法选择、数据输入、播放控制与复杂度信息面板，中央为可视化画布，右侧为代码面板与步骤面板。')
para('视觉风格采用“大理石奢华（Marble Luxury）”主题：以暖白大理石色为背景，玄武岩黑为主色，金色为点缀，灰褐为次要文字，'
     '并采用 Inter / Cormorant Garamond / JetBrains Mono 字体组合，整体观感简洁高级。主题色集中定义在 style.css 的 CSS 变量中，便于统一调整。')
caption_fig('图 3-1　系统主界面布局')

# ====================== 四、系统实现 ======================
heading('四、系统实现', 1)

heading('4.1 开发环境与工具', 2)
para('开发与运行环境如下：Node.js 18 及以上；构建工具 Vite 6；框架 Vue 3.5 与状态管理 Pinia 2；单元测试框架 Vitest 2；'
     '代码编辑器 Visual Studio Code；可视化基于浏览器原生 Canvas 2D，音效基于 Web Audio API。常用命令为 npm run dev（开发）、'
     'npm run build（构建）、npm test（测试）。')

heading('4.2 关键技术实现', 2)

heading('4.2.1 步骤生成器设计', 3)
para('步骤生成器是系统的核心。每个算法被实现为一个纯函数，在执行算法逻辑的同时，将每一个关键时刻的状态快照压入步骤数组并返回。'
     '以冒泡排序为例，在每次比较与交换时分别生成一个带有 highlights 标记和说明文字的步骤：')
code('''export function bubbleSort(arr) {
  const a = [...arr]
  const steps = []
  const n = a.length
  const sorted = new Set()
  steps.push({ array: [...a], highlights: {}, description: `初始数组: [${a.join(', ')}]`, codeLine: 0 })
  for (let i = 0; i < n - 1; i++) {
    for (let j = 0; j < n - 1 - i; j++) {
      const hl = {}; sorted.forEach(s => hl[s] = 'sorted')
      hl[j] = 'compare'; hl[j + 1] = 'compare'
      steps.push({ array: [...a], highlights: hl, description: `比较 a[${j}] 和 a[${j+1}]`, codeLine: 2 })
      if (a[j] > a[j + 1]) {
        ;[a[j], a[j + 1]] = [a[j + 1], a[j]]
        const hl2 = {}; sorted.forEach(s => hl2[s] = 'sorted')
        hl2[j] = 'swap'; hl2[j + 1] = 'swap'
        steps.push({ array: [...a], highlights: hl2, description: `交换`, codeLine: 3 })
      }
    }
    sorted.add(n - 1 - i)
  }
  return steps
}''', '代码 4-1　冒泡排序的步骤生成（节选）')

heading('4.2.2 Canvas 自适应渲染', 3)
para('为在不同尺寸与高分屏下保持清晰，画布在初始化与窗口变化时按设备像素比缩放，并据此重绘。渲染主函数依据步骤类型分发到具体绘制逻辑：')
code('''function resizeCanvas() {
  const rect = wrapperRef.value.getBoundingClientRect()
  const dpr = window.devicePixelRatio || 1
  canvasRef.value.width = rect.width * dpr
  canvasRef.value.height = rect.height * dpr
  canvasRef.value.style.width = rect.width + 'px'
  canvasRef.value.style.height = rect.height + 'px'
  ctx = canvasRef.value.getContext('2d')
  ctx.scale(dpr, dpr)
  render()
}

function render() {
  if (!ctx || !props.step) return
  if (props.step.type === 'hanoi') renderHanoi(ctx, w, h, props.step)
  else if (props.step.type === 'tree') renderTree(ctx, w, h, props.step)
  else if (props.step.array !== undefined) renderSortBars(ctx, w, h, props.step)
  // ……其余类型分发
}''', '代码 4-2　Canvas 高清适配与渲染分发（节选）')
para('对于汉诺塔等图形，还根据实际盘子数量与画布高度动态计算盘高、柱高并使整组图形垂直居中，避免在不同输入与窗口下出现溢出或偏底的问题。')

heading('4.2.3 响应式状态管理', 3)
para('系统使用 Pinia 组合式 Store 管理全局状态。stepInfo 为派生的计算属性，随 currentStep 自动更新；画布组件侦听步骤变化触发重绘，'
     '代码面板通过绑定 codeLine 高亮当前行，三者借助 Vue 的响应式系统自动联动：')
code('''const stepInfo = computed(() => {
  if (steps.value.length === 0) return null
  return steps.value[currentStep.value] || null
})

function stepForward() {
  if (currentStep.value < steps.value.length - 1) {
    currentStep.value++
    playStepSound(steps.value[currentStep.value])
    if (currentStep.value >= steps.value.length - 1) {
      playState.value = 'completed'; stopTimer()
      soundService.playCompleteSound()
    }
  }
}''', '代码 4-3　Pinia 状态管理与步骤推进（节选）')

heading('4.3 核心功能实现', 2)
para('播放控制由 Store 的定时器驱动：play 启动 setInterval，每个周期推进一步并播放音效，到达末尾自动停止；pause/togglePlay 控制暂停与继续；'
     'setSpeed 在播放中重启定时器以即时调整速度。代码同步通过将 stepInfo.codeLine 传入代码面板，由其高亮对应行实现。')
para('音效反馈由 soundService 统一合成。其根据步骤类型选取频率与波形，并通过增益包络与低通滤波得到柔和的音色，'
     '在比较、交换、访问、完成等不同操作上给予差异化的听觉提示。')

heading('4.4 AI 辅助开发说明', 2)
para('本项目在开发与文档撰写过程中借助了 AI 编程助手（Claude Code），主要用于以下方面：辅助统一界面主题配色与样式、'
     '协助补全与修正单元测试、生成与完善项目文档、对部分实现细节给出参考思路。')
para('所有 AI 生成或建议的代码与文字均经过团队成员的人工审阅、测试与修改，确保其与项目实际实现一致、符合课程设计要求。'
     'AI 仅作为提升效率的辅助工具，系统的设计思路、核心实现与最终成果均由团队成员自主完成并负责。')

# ====================== 五、系统测试 ======================
heading('五、系统测试', 1)

heading('5.1 测试环境', 2)
para('测试在以下环境进行：操作系统 Windows 11；运行环境 Node.js 18+；单元测试框架 Vitest 2.x；'
     '界面与动画在现代浏览器（Chrome / Edge）中进行手工验证。单元测试覆盖算法逻辑层的全部步骤生成函数与调度器。')

heading('5.2 功能测试', 2)
para('针对算法步骤生成与调度，编写了五个测试文件共 108 个用例，覆盖排序、图、数据结构、递归回溯与执行器，主要测试结果如表 5-1 所示。')
caption_fig('表 5-1　单元测试结果汇总')
tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_ALIGN.CENTER
hdr = tbl.rows[0].cells
for i, txt in enumerate(['测试文件', '测试对象', '用例数', '结果']):
    hdr[i].paragraphs[0].alignment = WD_ALIGN.CENTER
    rr = hdr[i].paragraphs[0].add_run(txt)
    set_font(rr, cn='黑体', en='SimHei', size=10.5, bold=True)
rows = [
    ['sorting.test.js', '6 种排序算法', '21', '通过'],
    ['graph.test.js', 'Dijkstra / BFS / DFS', '17', '通过'],
    ['dataStructures.test.js', '线性表 / 栈队列 / 树', '30', '通过'],
    ['recursion.test.js', '汉诺塔 / 棋盘 / 着色 / 迷宫', '23', '通过'],
    ['executor.test.js', '算法调度与随机数据', '17', '通过'],
    ['合计', '—', '108', '全部通过'],
]
for row in rows:
    cells = tbl.add_row().cells
    for i, txt in enumerate(row):
        cells[i].paragraphs[0].alignment = WD_ALIGN.CENTER
        rr = cells[i].paragraphs[0].add_run(txt)
        set_font(rr, cn='宋体', en='Times New Roman', size=10.5)
doc.add_paragraph()
para('功能测试要点包括：各排序算法对普通、已排序、逆序、含重复元素及负数数组均能正确排序，且不修改原数组；'
     '图算法对线性图、环形图、不连通图、零权与大权边均能正确生成遍历/最短路径步骤；'
     '递归回溯算法满足汉诺塔 2ⁿ−1 次移动、棋盘全覆盖、相邻节点异色、迷宫成功寻路等性质；调度器能正确分发各类算法并对未知 ID 返回空。')

heading('5.3 性能测试', 2)
para('对界面与构建性能进行验证：算法切换与单步操作响应及时，画布动画在常见数据规模下流畅无明显卡顿；'
     '播放速度可在 100~2000ms 间平滑调节；生产构建（npm run build）可正常完成，产物体积处于合理范围。')

heading('5.4 测试结果分析', 2)
para('全部 108 个单元测试用例均通过，表明算法逻辑层的步骤生成与调度功能正确可靠；界面与动画的手工测试结果符合预期。'
     '测试过程中也修正了早期测试与实现不一致的问题（如步骤字段命名），进一步保证了系统的稳定性与正确性。')

# ====================== 六、核心源代码 ======================
heading('六、核心源代码', 1)

heading('6.1 项目整体结构说明', 2)
para('项目采用标准的 Vite + Vue 工程结构，源代码集中在 src 目录，算法逻辑、状态管理、视图组件与工具服务分目录组织，'
     '测试位于 test 目录。整体结构如下：')
code('''算法可视化/
├── index.html              入口 HTML（引入字体）
├── vite.config.js          Vite 配置
├── vitest.config.js        测试配置
├── package.json
└── src/
    ├── main.js             应用入口
    ├── App.vue             主界面（三栏布局）
    ├── style.css           全局样式与主题变量
    ├── algorithms/         算法逻辑层
    │   ├── registry.js     算法注册表
    │   ├── executor.js     调度器与随机数据
    │   ├── sorting.js      排序算法
    │   ├── graph.js        图算法
    │   ├── recursion.js    递归与回溯
    │   └── dataStructures.js  线性表/栈队列/树
    ├── store/algorithmStore.js  Pinia 状态管理
    ├── utils/soundService.js    音效服务
    └── components/
        ├── VisualizationCanvas.vue  可视化画布
        ├── CodePanel.vue            代码面板
        └── StepsPanel.vue           步骤面板''', '代码 6-1　项目目录结构')

heading('6.2 核心模块代码说明', 2)

heading('6.2.1 main.js（应用入口）', 3)
para('应用入口创建 Vue 实例，注册 Pinia 并挂载根组件：')
code('''import { createApp } from 'vue'
import { createPinia } from 'pinia'
import App from './App.vue'
import './style.css'

const app = createApp(App)
app.use(createPinia())
app.mount('#app')''', '代码 6-2　main.js')

heading('6.2.2 algorithmStore.js（状态管理）', 3)
para('Store 使用组合式 API 定义全局状态与动作，集中管理算法、步骤序列与播放控制：')
code('''export const useAlgorithmStore = defineStore('algorithm', () => {
  const currentAlgoId = ref(null)
  const playState = ref('idle')       // idle/running/paused/completed
  const currentStep = ref(0)
  const steps = ref([])
  const playSpeed = ref(700)
  const soundEnabled = ref(true)
  let timer = null

  const totalSteps = computed(() => steps.value.length)
  const stepInfo = computed(() =>
    steps.value.length ? steps.value[currentStep.value] : null)

  function setSteps(newSteps) {
    steps.value = newSteps; currentStep.value = 0; playState.value = 'idle'
  }
  function play() { playState.value = 'running'; startTimer() }
  function pause() { playState.value = 'paused'; stopTimer() }
  function togglePlay() { playState.value === 'running' ? pause() : play() }

  function startTimer() {
    stopTimer()
    timer = setInterval(() => {
      if (currentStep.value < steps.value.length - 1) {
        currentStep.value++
        playStepSound(steps.value[currentStep.value])
      } else { playState.value = 'completed'; stopTimer() }
    }, playSpeed.value)
  }
  // ……stepForward / stepBack / reset / setSpeed / 音效控制
})''', '代码 6-3　algorithmStore.js（节选）')

heading('6.2.3 registry.js（算法注册表）', 3)
para('注册表以对象形式保存每个算法的元数据（名称、分类、复杂度、默认数据与源码），并定义分类列表：')
code('''export const algorithmRegistry = {
  'bubble-sort': {
    id: 'bubble-sort', name: '冒泡排序', category: 'sort',
    difficulty: '简单', timeComplexity: 'O(n²)', spaceComplexity: 'O(1)',
    description: '相邻元素两两比较，每轮将最大元素冒泡到末尾',
    inputType: 'array', defaultData: [64, 34, 25, 12, 22, 11, 90],
    code: `void BubbleSort(int a[], int n) { ... }`
  },
  // ……其余 31 个算法
}

export const categories = [
  { id: 'linear-list', name: '线性表', icon: '≡' },
  { id: 'stack-queue', name: '栈和队列', icon: '⊞' },
  { id: 'tree', name: '树结构', icon: '♧' },
  { id: 'graph', name: '图结构', icon: '◇' },
  { id: 'sort', name: '排序', icon: '↕' },
  { id: 'recursion', name: '递归与回溯', icon: '↻' }
]''', '代码 6-4　registry.js（节选）')

heading('6.2.4 executor.js（调度器）', 3)
para('调度器根据算法 ID 将输入数据分发到对应的步骤生成函数，是连接界面与算法逻辑层的桥梁：')
code('''export function executeAlgorithm(algoId, inputData) {
  const sortAlgos = { 'bubble-sort': bubbleSort, 'quick-sort': quickSort, /* … */ }
  const graphAlgos = { 'dijkstra': dijkstra, 'bfs': bfs, 'dfs': dfs }

  if (sortAlgos[algoId]) return sortAlgos[algoId](inputData)
  if (graphAlgos[algoId]) return graphAlgos[algoId](inputData)
  if (algoId === 'hanoi') return hanoi(inputData)
  if (algoId === 'maze') return mazeSolve(inputData)
  if (linearListAlgos.includes(algoId)) return linearListSteps(algoId, inputData)
  if (stackQueueAlgos.includes(algoId)) return stackQueueSteps(algoId, inputData)
  if (treeAlgos.includes(algoId)) return treeSteps(algoId, inputData)
  return []
}''', '代码 6-5　executor.js（节选）')

heading('6.2.5 sorting.js（排序算法）', 3)
para('排序模块实现了冒泡、快速、插入、希尔、选择、归并六种排序，均以纯函数返回步骤序列，互不依赖界面。其冒泡排序实现见代码 4-1。')

heading('6.3 核心组件实现', 2)

heading('6.3.1 VisualizationCanvas.vue（可视化画布）', 3)
para('画布组件接收当前算法 ID 与步骤对象，依据类型绘制对应图形。排序柱状图的核心绘制如下：')
code('''function renderSortBars(ctx, w, h, step) {
  const arr = step.array
  const maxVal = Math.max(...arr, 1)
  const barW = Math.max(10, (w - 120 - gap * (arr.length - 1)) / arr.length)
  arr.forEach((val, i) => {
    const barH = (val / maxVal) * barAreaH
    const hl = step.highlights?.[i]
    ctx.fillStyle = hl ? COLORS[hl] || COLORS.default : COLORS.default
    ctx.roundRect(x, y, barW, barH, [4, 4, 0, 0]); ctx.fill()
    ctx.fillStyle = '#1A1A1A'           // 值标签
    ctx.fillText(val, x + barW / 2, y - 8)
  })
}''', '代码 6-6　VisualizationCanvas.vue 排序渲染（节选）')

heading('6.3.2 CodePanel.vue（代码高亮）', 3)
para('代码面板将源码按行展示，并根据传入的 highlightLine 为当前执行行添加高亮样式：')
code('''<div class="code-body">
  <div v-for="(line, i) in codeLines" :key="i"
       class="code-line" :class="{ highlighted: i === highlightLine }">
    <span class="marker">{{ i === highlightLine ? '▶' : '' }}</span>
    <span class="line-num">{{ i + 1 }}</span>
    <span class="line-content">{{ line }}</span>
  </div>
</div>''', '代码 6-7　CodePanel.vue 模板（节选）')

heading('6.3.3 StepsPanel.vue（步骤面板）', 3)
para('步骤面板列出全部步骤，高亮当前步骤并支持点击跳转：')
code('''<div class="steps-body">
  <div v-for="(step, index) in store.steps" :key="index"
       class="step-row" :class="{ active: index === store.currentStep }"
       @click="goto(index)">
    <span class="step-no">#{{ index + 1 }}</span>
    <span class="step-title">{{ step.title || '步骤' }}</span>
    <span class="step-desc">{{ step.description }}</span>
  </div>
</div>''', '代码 6-8　StepsPanel.vue 模板（节选）')

heading('6.4 界面与样式设计', 2)
para('全局样式与主题色集中定义在 style.css 的根变量中，采用“大理石奢华”配色，修改这些变量即可统一切换全站风格：')
code(''':root {
  --bg-dark: #F8F6F3;       /* 大理石暖白背景 */
  --bg-card: #FFFFFF;
  --primary: #1A1A1A;       /* 玄武岩黑（主色） */
  --accent: #C9A96E;        /* 金色点缀 */
  --text-primary: #1A1A1A;
  --text-secondary: #8A7968;/* 灰褐文字 */
  --success: #85BD73;       /* 绿色 */
  --danger: #C0564B;        /* 赤陶红 */
  --border: #E8E0D6;
}''', '代码 6-9　style.css 主题变量（节选）')

# ---------- 保存 ----------
out = '算法过程可视化系统_课程设计报告.docx'
doc.save(out)
print('SAVED:', out)
